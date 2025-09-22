from docx import Document

doc = Document()
doc.add_heading('Решение всех заданий по homework_08', level=1)

# -------------------- 1. Линейная алгебра --------------------
doc.add_heading('1. Линейная алгебра', level=2)

doc.add_heading('1) Найти D = A^T C - 2 A^T B^T', level=3)
doc.add_paragraph("Ответ: D = [[14, 14, -11], [-16, 3, -6]]")

doc.add_heading('2) Решение системы на элементы матриц', level=3)
doc.add_paragraph("Ответ: x = 2, y = 6, z = -4, v = 10")

doc.add_heading('3) Ранг матрицы 1', level=3)
doc.add_paragraph("Ответ: p = 9, q = 15")

doc.add_heading('4) Смена базиса', level=3)
doc.add_paragraph("Ответ: [x]_B = (-1, -3),   y = (1, -2)")

doc.add_heading('5) Низкоранговая аппроксимация SVD', level=3)
doc.add_paragraph(
    "График ошибки ||A - A_r||_F убывает с ростом r; точка излома определяет оптимальный r."
)

# -------------------- 2. Математический анализ и оптимизация --------------------
doc.add_heading('2. Математический анализ и оптимизация', level=2)

doc.add_heading('1) Градиент и критические точки', level=3)
doc.add_paragraph(
    "Градиент: (3x1^2 - 2x2 - 3,  -2x1 + 2x2 - 2)\n"
    "Критические точки: (-1, 0),  (5/3, 8/3)"
)

doc.add_heading('2) Проверка уравнения для f = ln(sqrt(x1)+sqrt(x2))', level=3)
doc.add_paragraph(
    "Равенство x1*∂f/∂x1 + x2*∂f/∂x2 = 1/2 выполняется для x1, x2 > 0."
)

doc.add_heading('3) Матрица Якоби', level=3)
doc.add_paragraph(
    "J_f(x, y, z) = [[1, 1, 1], [yz, xz, xy]]\n"
    "В точке (1, 2, 3): [[1, 1, 1], [6, 3, 2]]"
)

doc.add_heading('4) f(x) = 1/3 ||x||^3', level=3)
doc.add_paragraph("Градиент: ||x|| * x")

doc.add_heading('5) f(x) = ||x||', level=3)
doc.add_paragraph("Градиент: x / ||x||")

doc.add_heading('6) f(x) = ||A x||^2', level=3)
doc.add_paragraph("Градиент: 2 A^T A x")

doc.add_heading('7) f(x) = -e^{-x^T x}', level=3)
doc.add_paragraph("Градиент: 2 e^{-x^T x} x")

# Сохраняем документ
doc.save('homework_08_solutions.docx')

print("Файл homework_08_solutions.docx успешно создан.")
